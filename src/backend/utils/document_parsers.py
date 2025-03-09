import os
import logging
from abc import ABC, abstractmethod
from typing import List, Dict, Optional, Any, Union

import fitz  # PyMuPDF v1.23.0
import docx  # python-docx v0.8.11
from docx import Document
import pandas as pd  # pandas v2.0.0

from ..schemas.document import ALLOWED_FILE_TYPES, DocumentChunk
from ..utils.text_processing import TextChunker, clean_text
from ..config.settings import Settings

# Configure logger
logger = logging.getLogger(__name__)

# Initialize settings
settings = Settings()

# Default constants
DEFAULT_CHUNK_SIZE = 1000
DEFAULT_CHUNK_OVERLAP = 100


class DocumentParser(ABC):
    """Abstract base class defining the interface for document parsers"""
    
    def __init__(self, options: Optional[Dict[str, Any]] = None):
        """Initializes the document parser"""
        self.options = options or {}
        self.chunk_size = self.options.get('chunk_size', DEFAULT_CHUNK_SIZE)
        self.chunk_overlap = self.options.get('chunk_overlap', DEFAULT_CHUNK_OVERLAP)
        self.text_chunker = TextChunker(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap
        )
    
    @abstractmethod
    def parse_file(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> List[DocumentChunk]:
        """
        Parses a document file and extracts its content as chunks
        
        Args:
            file_path: Path to the document file
            metadata: Optional metadata to include with chunks
            
        Returns:
            List of document chunks with extracted content
        """
        pass
    
    @abstractmethod
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extracts document-specific metadata
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing document metadata
        """
        pass
    
    def create_chunks(self, document_id: str, text_chunks: List[str], 
                     metadata: Optional[Dict[str, Any]] = None,
                     page_numbers: Optional[List[int]] = None) -> List[DocumentChunk]:
        """
        Creates DocumentChunk objects from extracted text
        
        Args:
            document_id: ID of the document
            text_chunks: List of text chunks
            metadata: Optional metadata to include with chunks
            page_numbers: Optional list of page numbers for each chunk
            
        Returns:
            List of document chunks
        """
        chunks = []
        for i, chunk in enumerate(text_chunks):
            chunk_data = {
                "document_id": document_id,
                "chunk_index": i,
                "content": chunk,
            }
            
            if page_numbers and i < len(page_numbers):
                chunk_data["page_number"] = page_numbers[i]
                
            if metadata:
                chunk_data["metadata"] = metadata
                
            chunks.append(DocumentChunk(**chunk_data))
            
        return chunks


class PDFParser(DocumentParser):
    """Parser for PDF documents using PyMuPDF (fitz)"""
    
    def __init__(self, options: Optional[Dict[str, Any]] = None):
        """Initializes the PDF parser"""
        super().__init__(options)
        # Set PDF-specific options
        self.extract_images = self.options.get('extract_images', False)
        self.extract_tables = self.options.get('extract_tables', False)
    
    def parse_file(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> List[DocumentChunk]:
        """
        Parses a PDF file and extracts its content as chunks
        
        Args:
            file_path: Path to the PDF file
            metadata: Optional metadata to include with chunks
            
        Returns:
            List of document chunks with extracted content
        """
        try:
            # Open PDF file
            document = fitz.open(file_path)
            
            text_chunks = []
            page_numbers = []
            
            # Process each page
            for page_num, page in enumerate(document):
                # Extract text from page
                page_text = page.get_text()
                
                # Clean text
                cleaned_text = clean_text(page_text)
                
                if not cleaned_text:
                    continue
                
                # Split into chunks
                page_chunks = self.text_chunker.split_by_semantic_units(cleaned_text)
                
                # Add to overall chunks with page tracking
                text_chunks.extend(page_chunks)
                page_numbers.extend([page_num + 1] * len(page_chunks))
            
            document_id = metadata.get('document_id', '') if metadata else ''
            
            # Create document chunks
            return self.create_chunks(document_id, text_chunks, metadata, page_numbers)
            
        except Exception as e:
            logger.error(f"Error parsing PDF file {file_path}: {str(e)}")
            raise
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extracts PDF-specific metadata
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Dictionary containing PDF metadata
        """
        try:
            # Get basic file metadata
            file_metadata = extract_file_metadata(file_path)
            
            # Open PDF file
            document = fitz.open(file_path)
            
            # Extract PDF-specific metadata
            pdf_metadata = {
                'page_count': len(document),
                'pdf_version': document.pdf_version,
                'is_encrypted': document.is_encrypted,
                'pdf_metadata': document.metadata
            }
            
            # Combine all metadata
            return {**file_metadata, **pdf_metadata}
            
        except Exception as e:
            logger.error(f"Error extracting metadata from PDF file {file_path}: {str(e)}")
            return extract_file_metadata(file_path)  # Fall back to basic metadata


class DocxParser(DocumentParser):
    """Parser for Microsoft Word documents using python-docx"""
    
    def __init__(self, options: Optional[Dict[str, Any]] = None):
        """Initializes the Word document parser"""
        super().__init__(options)
        # Set DOCX-specific options
        self.extract_headers = self.options.get('extract_headers', True)
        self.extract_footers = self.options.get('extract_footers', True)
    
    def parse_file(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> List[DocumentChunk]:
        """
        Parses a Word document and extracts its content as chunks
        
        Args:
            file_path: Path to the Word document
            metadata: Optional metadata to include with chunks
            
        Returns:
            List of document chunks with extracted content
        """
        try:
            # Open DOCX file
            document = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
            full_text = '\n\n'.join(paragraphs)
            
            # Clean text
            cleaned_text = clean_text(full_text)
            
            # Split into chunks
            text_chunks = self.text_chunker.split_by_semantic_units(cleaned_text)
            
            document_id = metadata.get('document_id', '') if metadata else ''
            
            # Create document chunks
            return self.create_chunks(document_id, text_chunks, metadata)
            
        except Exception as e:
            logger.error(f"Error parsing Word document {file_path}: {str(e)}")
            raise
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extracts Word document-specific metadata
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            Dictionary containing Word document metadata
        """
        try:
            # Get basic file metadata
            file_metadata = extract_file_metadata(file_path)
            
            # Open DOCX file
            document = Document(file_path)
            
            # Extract document properties
            doc_metadata = {}
            if document.core_properties:
                props = document.core_properties
                doc_metadata = {
                    'author': props.author,
                    'title': props.title,
                    'created': props.created,
                    'modified': props.modified,
                    'paragraph_count': len(document.paragraphs)
                }
            
            # Combine all metadata
            return {**file_metadata, **doc_metadata}
            
        except Exception as e:
            logger.error(f"Error extracting metadata from Word document {file_path}: {str(e)}")
            return extract_file_metadata(file_path)  # Fall back to basic metadata


class TextParser(DocumentParser):
    """Parser for plain text files including TXT and Markdown"""
    
    def __init__(self, options: Optional[Dict[str, Any]] = None):
        """Initializes the text file parser"""
        super().__init__(options)
        # Set text-specific options
        self.encoding = self.options.get('encoding', 'utf-8')
    
    def parse_file(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> List[DocumentChunk]:
        """
        Parses a text file and extracts its content as chunks
        
        Args:
            file_path: Path to the text file
            metadata: Optional metadata to include with chunks
            
        Returns:
            List of document chunks with extracted content
        """
        try:
            # Try to read with specified encoding
            try:
                with open(file_path, 'r', encoding=self.encoding) as f:
                    content = f.read()
            except UnicodeDecodeError:
                # Fall back to latin-1 which should handle any byte sequence
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()
            
            # Clean text
            cleaned_text = clean_text(content)
            
            # Split into chunks
            text_chunks = self.text_chunker.split_by_semantic_units(cleaned_text)
            
            document_id = metadata.get('document_id', '') if metadata else ''
            
            # Create document chunks
            return self.create_chunks(document_id, text_chunks, metadata)
            
        except Exception as e:
            logger.error(f"Error parsing text file {file_path}: {str(e)}")
            raise
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extracts text file metadata
        
        Args:
            file_path: Path to the text file
            
        Returns:
            Dictionary containing text file metadata
        """
        try:
            # Get basic file metadata
            file_metadata = extract_file_metadata(file_path)
            
            # Determine if file is markdown based on extension
            is_markdown = file_path.lower().endswith('.md')
            
            # Count lines in the file
            line_count = 0
            word_count = 0
            
            try:
                with open(file_path, 'r', encoding=self.encoding) as f:
                    for line in f:
                        line_count += 1
                        word_count += len(line.split())
            except UnicodeDecodeError:
                with open(file_path, 'r', encoding='latin-1') as f:
                    for line in f:
                        line_count += 1
                        word_count += len(line.split())
            
            # Add text-specific metadata
            text_metadata = {
                'line_count': line_count,
                'word_count': word_count,
                'is_markdown': is_markdown
            }
            
            # Combine all metadata
            return {**file_metadata, **text_metadata}
            
        except Exception as e:
            logger.error(f"Error extracting metadata from text file {file_path}: {str(e)}")
            return extract_file_metadata(file_path)  # Fall back to basic metadata


class SpreadsheetParser(DocumentParser):
    """Parser for spreadsheet files including CSV and Excel using pandas"""
    
    def __init__(self, options: Optional[Dict[str, Any]] = None):
        """Initializes the spreadsheet parser"""
        super().__init__(options)
        # Set spreadsheet-specific options
        self.sheet_names = self.options.get('sheet_names', None)  # None means all sheets
        self.max_rows = self.options.get('max_rows', None)  # None means all rows
    
    def parse_file(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> List[DocumentChunk]:
        """
        Parses a spreadsheet file and extracts its content as chunks
        
        Args:
            file_path: Path to the spreadsheet file
            metadata: Optional metadata to include with chunks
            
        Returns:
            List of document chunks with extracted content
        """
        try:
            # Determine file type based on extension
            file_extension = os.path.splitext(file_path)[1].lower()
            
            # Read the spreadsheet
            if file_extension == '.csv':
                df = pd.read_csv(file_path, nrows=self.max_rows)
                text_content = self.dataframe_to_text(df)
            else:  # Excel files
                if self.sheet_names:
                    # Read specific sheets
                    dfs = []
                    for sheet in self.sheet_names:
                        dfs.append(pd.read_excel(file_path, sheet_name=sheet, nrows=self.max_rows))
                    
                    # Convert each sheet to text and join with clear separators
                    sheet_texts = []
                    for i, df in enumerate(dfs):
                        sheet_name = self.sheet_names[i]
                        sheet_text = f"--- Sheet: {sheet_name} ---\n{self.dataframe_to_text(df)}"
                        sheet_texts.append(sheet_text)
                    
                    text_content = "\n\n".join(sheet_texts)
                else:
                    # Read all sheets
                    excel_data = pd.read_excel(file_path, sheet_name=None, nrows=self.max_rows)
                    
                    # Convert each sheet to text and join with clear separators
                    sheet_texts = []
                    for sheet_name, df in excel_data.items():
                        sheet_text = f"--- Sheet: {sheet_name} ---\n{self.dataframe_to_text(df)}"
                        sheet_texts.append(sheet_text)
                    
                    text_content = "\n\n".join(sheet_texts)
            
            # Clean text
            cleaned_text = clean_text(text_content)
            
            # Split into chunks
            text_chunks = self.text_chunker.split_by_semantic_units(cleaned_text)
            
            document_id = metadata.get('document_id', '') if metadata else ''
            
            # Create document chunks
            return self.create_chunks(document_id, text_chunks, metadata)
            
        except Exception as e:
            logger.error(f"Error parsing spreadsheet file {file_path}: {str(e)}")
            raise
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extracts spreadsheet metadata
        
        Args:
            file_path: Path to the spreadsheet file
            
        Returns:
            Dictionary containing spreadsheet metadata
        """
        try:
            # Get basic file metadata
            file_metadata = extract_file_metadata(file_path)
            
            # Determine file type based on extension
            file_extension = os.path.splitext(file_path)[1].lower()
            
            # Extract spreadsheet-specific metadata
            spreadsheet_metadata = {}
            
            if file_extension == '.csv':
                # Read CSV file
                df = pd.read_csv(file_path)
                spreadsheet_metadata = {
                    'row_count': len(df),
                    'column_count': len(df.columns),
                    'columns': list(df.columns),
                    'file_type': 'csv'
                }
            else:  # Excel files
                # Read Excel file info
                excel_data = pd.read_excel(file_path, sheet_name=None)
                sheet_names = list(excel_data.keys())
                
                # Count total rows across all sheets
                total_rows = sum(len(df) for df in excel_data.values())
                
                # Get column counts for each sheet
                sheet_info = {
                    sheet: {
                        'rows': len(df),
                        'columns': len(df.columns)
                    }
                    for sheet, df in excel_data.items()
                }
                
                spreadsheet_metadata = {
                    'sheet_count': len(sheet_names),
                    'sheet_names': sheet_names,
                    'total_row_count': total_rows,
                    'sheet_info': sheet_info,
                    'file_type': 'excel'
                }
            
            # Combine all metadata
            return {**file_metadata, **spreadsheet_metadata}
            
        except Exception as e:
            logger.error(f"Error extracting metadata from spreadsheet file {file_path}: {str(e)}")
            return extract_file_metadata(file_path)  # Fall back to basic metadata
    
    def dataframe_to_text(self, df: pd.DataFrame) -> str:
        """
        Converts a pandas DataFrame to a text representation
        
        Args:
            df: DataFrame to convert
            
        Returns:
            Text representation of the dataframe
        """
        try:
            # Convert DataFrame to string representation
            text = df.to_string(index=True, header=True)
            
            # Clean and normalize the text
            return clean_text(text)
        except Exception as e:
            logger.error(f"Error converting dataframe to text: {str(e)}")
            return str(df)  # Fall back to basic string representation


def get_parser_for_file_type(file_type: str) -> DocumentParser:
    """
    Factory function that returns the appropriate document parser for a given file type
    
    Args:
        file_type: File extension or type
        
    Returns:
        An instance of the appropriate parser for the file type
        
    Raises:
        ValueError: If the file type is not supported
    """
    # Check if file type is supported
    if file_type.lower() not in ALLOWED_FILE_TYPES:
        raise ValueError(f"Unsupported file type: {file_type}. Supported types: {', '.join(ALLOWED_FILE_TYPES)}")
    
    # Return the appropriate parser based on file type
    file_type = file_type.lower()
    
    if file_type == 'pdf':
        return PDFParser()
    elif file_type == 'docx':
        return DocxParser()
    elif file_type in ['txt', 'md']:
        return TextParser()
    elif file_type in ['csv', 'xlsx']:
        return SpreadsheetParser()
    else:
        raise ValueError(f"No parser available for file type: {file_type}")


def extract_file_metadata(file_path: str) -> Dict[str, Any]:
    """
    Extracts basic metadata from a file such as size, creation date, and modification date
    
    Args:
        file_path: Path to the file
        
    Returns:
        Dictionary containing file metadata
    """
    try:
        # Check if file exists
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Get file stats
        file_stats = os.stat(file_path)
        
        # Extract basic metadata
        metadata = {
            'file_name': os.path.basename(file_path),
            'file_extension': os.path.splitext(file_path)[1].lower().lstrip('.'),
            'file_size_bytes': file_stats.st_size,
            'file_size_kb': round(file_stats.st_size / 1024, 2),
            'creation_time': file_stats.st_ctime,
            'modification_time': file_stats.st_mtime,
            'file_path': file_path
        }
        
        return metadata
    except Exception as e:
        logger.error(f"Error extracting file metadata for {file_path}: {str(e)}")
        # Return minimal metadata on error
        return {
            'file_name': os.path.basename(file_path),
            'file_path': file_path,
            'error': str(e)
        }